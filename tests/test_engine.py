# Copyright (C) 2026 Carota-Bunny
# SPDX-License-Identifier: AGPL-3.0-only

from __future__ import annotations

import tempfile
import unittest
from datetime import datetime
from pathlib import Path
from unittest.mock import patch
from zipfile import ZIP_DEFLATED, ZIP_STORED, ZipFile, ZipInfo
import xml.etree.ElementTree as ET

import fitz
from docx import Document
from PIL import Image, ImageCms
from PIL.PngImagePlugin import PngInfo

from metacleaner.engine import (
    _serialize_with_default_namespace,
    clean_file,
    inspect_file,
    unique_output_path,
)
from metacleaner.models import CleanMode, CleanOptions


REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def rewrite_zip(path: Path, updates: dict[str, bytes], removals: set[str] | None = None) -> None:
    removals = removals or set()
    with ZipFile(path, "r") as zin:
        entries = [(info, zin.read(info.filename)) for info in zin.infolist()]
    temp = path.with_suffix(path.suffix + ".tmp")
    with ZipFile(temp, "w", compression=ZIP_DEFLATED) as zout:
        seen: set[str] = set()
        for info, data in entries:
            if info.filename in removals:
                continue
            if info.filename == "word/document.xml":
                info.date_time = (2020, 2, 3, 4, 5, 6)
                info.comment = b"Alice-comment"
            if info.filename in updates:
                data = updates[info.filename]
            zout.writestr(info, data)
            seen.add(info.filename)
        for name, data in updates.items():
            if name not in seen:
                zout.writestr(name, data)
    temp.replace(path)


def make_docx(path: Path, signed: bool = False) -> None:
    document = Document()
    document.add_paragraph("正文内容不得改变")
    props = document.core_properties
    props.author = "Alice Example"
    props.last_modified_by = "Bob Example"
    props.title = "保留的文档标题"
    props.subject = "保留的主题"
    props.keywords = "uranium, membrane"
    props.created = datetime(2021, 12, 20, 12, 14, 0)
    props.modified = datetime(2026, 6, 30, 5, 51, 0)
    props.last_printed = datetime(2026, 6, 29, 8, 30, 0)
    document.save(path)

    with ZipFile(path, "r") as zf:
        rels = ET.fromstring(zf.read("_rels/.rels"))
        types = ET.fromstring(zf.read("[Content_Types].xml"))
        app = ET.fromstring(zf.read("docProps/app.xml"))

    next_rid = f"rId{len(rels) + 10}"
    ET.SubElement(
        rels,
        f"{{{REL_NS}}}Relationship",
        {
            "Id": next_rid,
            "Type": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/custom-properties",
            "Target": "private/metadata.xml",
        },
    )
    ET.SubElement(
        rels,
        f"{{{REL_NS}}}Relationship",
        {
            "Id": next_rid + "Thumb",
            "Type": "http://schemas.openxmlformats.org/package/2006/relationships/metadata/thumbnail",
            "Target": "docProps/thumbnail.jpeg",
        },
    )
    ET.SubElement(
        types,
        f"{{{CT_NS}}}Override",
        {
            "PartName": "/private/metadata.xml",
            "ContentType": "application/vnd.openxmlformats-officedocument.custom-properties+xml",
        },
    )
    ET.SubElement(
        types,
        f"{{{CT_NS}}}Override",
        {"PartName": "/docProps/thumbnail.jpeg", "ContentType": "image/jpeg"},
    )
    app_namespace = app.tag.split("}", 1)[0].lstrip("{")
    app_values = {
        "Manager": "Sensitive Manager",
        "Company": "Sensitive Laboratory",
        "Template": "C:\\Users\\Alice\\AppData\\Roaming\\Microsoft\\Templates\\Normal.dotm",
        "HyperlinkBase": "C:\\Users\\Alice\\Documents",
    }
    existing_app = {elem.tag.rsplit("}", 1)[-1]: elem for elem in app}
    for name, value in app_values.items():
        elem = existing_app.get(name)
        if elem is None:
            elem = ET.SubElement(app, f"{{{app_namespace}}}{name}")
        elem.text = value

    custom = b"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/custom-properties"
 xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
 <property fmtid="{D5CDD505-2E9C-101B-9397-08002B2CF9AE}" pid="2" name="EmployeeId"><vt:lpwstr>EMP-001</vt:lpwstr></property>
</Properties>"""
    comments = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="{W_NS}">
 <w:comment w:id="0" w:author="Reviewer Name" w:initials="RN" w:date="2026-07-21T01:02:03Z">
  <w:p><w:r><w:t>审阅意见</w:t></w:r></w:p>
 </w:comment>
</w:comments>""".encode("utf-8")
    header = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:hdr xmlns:w="{W_NS}"
 xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
 xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
 xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"
 mc:Ignorable="w14 w15"><w:p>
  <w:ins w:id="7" w:author="Header Reviewer" w:date="2026-07-20T01:02:03Z">
   <w:r><w:t>页眉修订</w:t></w:r>
  </w:ins>
  <w:moveFromRangeStart w:id="8" w:author="Move Reviewer" w:date="2026-07-19T01:02:03Z"/>
</w:p></w:hdr>""".encode("utf-8")
    people = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w15:people xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml">
 <w15:person w15:author="Reviewer Name" w15:id="{11111111-2222-3333-4444-555555555555}">
  <w15:presenceInfo w15:providerId="SensitiveProvider" w15:userId="reviewer@example.test"/>
 </w15:person>
</w15:people>""".encode("utf-8")
    threaded_comments = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w15:threadedComments xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml">
 <w15:threadedComment w15:id="{AAAAAAAA-BBBB-CCCC-DDDD-EEEEEEEEEEEE}"
  w15:personId="{11111111-2222-3333-4444-555555555555}"
  w15:dT="2026-07-18T01:02:03Z">批注正文保留</w15:threadedComment>
</w15:threadedComments>""".encode("utf-8")
    powerpoint_comment = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:cmLst xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">
 <p:cm authorId="7" dt="2026-07-17T01:02:03Z"><p:text>批注正文保留</p:text></p:cm>
</p:cmLst>""".encode("utf-8")
    updates = {
        "_rels/.rels": ET.tostring(rels, encoding="utf-8", xml_declaration=True),
        "[Content_Types].xml": ET.tostring(types, encoding="utf-8", xml_declaration=True),
        "docProps/app.xml": ET.tostring(app, encoding="utf-8", xml_declaration=True),
        "private/metadata.xml": custom,
        "docProps/thumbnail.jpeg": b"not-a-rendered-thumbnail",
        "word/comments.xml": comments,
        "word/header1.xml": header,
        "word/people.xml": people,
        "word/threadedComments.xml": threaded_comments,
        "ppt/comments/comment1.xml": powerpoint_comment,
    }
    if signed:
        updates["_xmlsignatures/sig1.xml"] = b"<Signature/>"
    rewrite_zip(path, updates)


def make_odt(path: Path) -> None:
    meta = """<?xml version="1.0" encoding="UTF-8"?>
<office:document-meta
 xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"
 xmlns:dc="http://purl.org/dc/elements/1.1/"
 xmlns:meta="urn:oasis:names:tc:opendocument:xmlns:meta:1.0"
 xmlns:xlink="http://www.w3.org/1999/xlink">
 <office:meta>
  <meta:initial-creator>Alice Example</meta:initial-creator>
  <dc:creator>Bob Example</dc:creator>
  <meta:creation-date>2026-07-21T01:02:03</meta:creation-date>
  <dc:title>保留的 ODF 标题</dc:title>
   <meta:user-defined meta:name="EmployeeId">EMP-001</meta:user-defined>
   <meta:template xlink:href="file:///C:/Users/Alice/Templates/private-template.ott"/>
 </office:meta>
</office:document-meta>""".encode("utf-8")
    content = """<?xml version="1.0" encoding="UTF-8"?>
<office:document-content
 xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"
 xmlns:dc="http://purl.org/dc/elements/1.1/"
 xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0">
 <office:body><office:text><text:p>正文内容不得改变：<text:date>2026-08-15</text:date></text:p>
 <text:changed-region><office:change-info><dc:creator>Reviewer Name</dc:creator><dc:date>2026-07-21</dc:date></office:change-info></text:changed-region>
 </office:text></office:body>
</office:document-content>""".encode("utf-8")
    manifest = b"""<?xml version="1.0" encoding="UTF-8"?>
<manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0">
 <manifest:file-entry manifest:full-path="/" manifest:media-type="application/vnd.oasis.opendocument.text"/>
 <manifest:file-entry manifest:full-path="content.xml" manifest:media-type="text/xml"/>
 <manifest:file-entry manifest:full-path="meta.xml" manifest:media-type="text/xml"/>
</manifest:manifest>"""
    with ZipFile(path, "w") as zf:
        info = ZipInfo("mimetype")
        info.compress_type = ZIP_STORED
        zf.writestr(info, b"application/vnd.oasis.opendocument.text")
        zf.writestr("content.xml", content, compress_type=ZIP_DEFLATED)
        zf.writestr("meta.xml", meta, compress_type=ZIP_DEFLATED)
        zf.writestr("META-INF/manifest.xml", manifest, compress_type=ZIP_DEFLATED)


class EngineTests(unittest.TestCase):
    def setUp(self) -> None:
        self.tmp = tempfile.TemporaryDirectory()
        self.root = Path(self.tmp.name)

    def tearDown(self) -> None:
        self.tmp.cleanup()

    def test_docx_personal_mode_preserves_content_and_descriptive_properties(self) -> None:
        source = self.root / "sample.docx"
        make_docx(source)
        before = inspect_file(source)
        self.assertTrue(before.supported)
        self.assertGreaterEqual(before.personal_count, 5)

        result = clean_file(source, options=CleanOptions(mode=CleanMode.PERSONAL))
        self.assertTrue(result.success, result.error or result.residual_items)
        self.assertIsNotNone(result.output)
        output = result.output
        assert output is not None
        self.assertNotEqual(source, output)
        self.assertEqual(Document(output).paragraphs[0].text, "正文内容不得改变")

        props = Document(output).core_properties
        self.assertFalse(props.author)
        self.assertFalse(props.last_modified_by)
        self.assertEqual(props.title, "保留的文档标题")
        self.assertEqual(props.subject, "保留的主题")
        with ZipFile(output) as zf:
            names = set(zf.namelist())
            self.assertNotIn("docProps/custom.xml", names)
            self.assertNotIn("private/metadata.xml", names)
            self.assertNotIn("docProps/thumbnail.jpeg", names)
            content_types_xml = zf.read("[Content_Types].xml")
            root_relationships_xml = zf.read("_rels/.rels")
            app_xml = zf.read("docProps/app.xml")
            self.assertIn(b'<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"', content_types_xml)
            self.assertNotIn(b"<ct:Types", content_types_xml)
            self.assertIn(b'<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"', root_relationships_xml)
            self.assertNotIn(b"<rel:Relationships", root_relationships_xml)
            self.assertIn(
                b'<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"',
                app_xml,
            )
            comments = zf.read("word/comments.xml").decode("utf-8")
            header = zf.read("word/header1.xml").decode("utf-8")
            self.assertIn("Anonymous", comments)
            self.assertNotIn("Reviewer Name", comments)
            self.assertNotIn("2026-07-21T01:02:03Z", comments)
            self.assertIn("Anonymous", header)
            self.assertNotIn("Header Reviewer", header)
            self.assertNotIn("2026-07-20T01:02:03Z", header)
            self.assertNotIn("Move Reviewer", header)
            self.assertNotIn("2026-07-19T01:02:03Z", header)
            self.assertIn('mc:Ignorable="w14 w15"', header)
            self.assertIn('xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"', header)
            self.assertIn('xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"', header)
            ET.fromstring(header.encode("utf-8"))
            people = zf.read("word/people.xml").decode("utf-8")
            self.assertNotIn("reviewer@example.test", people)
            self.assertNotIn("SensitiveProvider", people)
            self.assertNotIn("11111111-2222-3333-4444-555555555555", people)
            threaded = zf.read("word/threadedComments.xml").decode("utf-8")
            self.assertNotIn("11111111-2222-3333-4444-555555555555", threaded)
            self.assertNotIn("2026-07-18T01:02:03Z", threaded)
            self.assertIn("批注正文保留", threaded)
            powerpoint = zf.read("ppt/comments/comment1.xml").decode("utf-8")
            self.assertNotIn("2026-07-17T01:02:03Z", powerpoint)
            self.assertIn('authorId="7"', powerpoint)
            self.assertIn("批注正文保留", powerpoint)
            core = ET.fromstring(zf.read("docProps/core.xml"))
            core_names = {elem.tag.rsplit("}", 1)[-1] for elem in core}
            self.assertTrue(
                {"creator", "lastModifiedBy", "created", "modified", "lastPrinted"}.isdisjoint(core_names),
                core_names,
            )
            app = ET.fromstring(app_xml)
            app_names = {elem.tag.rsplit("}", 1)[-1] for elem in app}
            self.assertTrue({"Manager", "Company", "Template", "HyperlinkBase"}.isdisjoint(app_names), app_names)
            self.assertTrue({"Application", "AppVersion"}.issubset(app_names), app_names)
            for info in zf.infolist():
                self.assertEqual(info.date_time, (1980, 1, 1, 0, 0, 0))
                self.assertEqual(info.comment, b"")
                self.assertEqual(info.extra, b"")
        after = inspect_file(output)
        self.assertEqual(after.personal_count, 0, after.items)

    def test_docx_never_contains_empty_typed_w3cdtf_dates(self) -> None:
        source = self.root / "typed-dates.docx"
        make_docx(source)
        result = clean_file(source, options=CleanOptions(mode=CleanMode.PERSONAL))
        self.assertTrue(result.success, result.error or result.residual_items)
        assert result.output is not None
        with ZipFile(result.output) as zf:
            core = ET.fromstring(zf.read("docProps/core.xml"))
            for elem in core:
                xsi_type = elem.attrib.get("{http://www.w3.org/2001/XMLSchema-instance}type")
                self.assertFalse(xsi_type == "dcterms:W3CDTF" and not (elem.text or "").strip())

    def test_docx_all_mode_removes_property_parts(self) -> None:
        source = self.root / "all.docx"
        make_docx(source)
        result = clean_file(source, options=CleanOptions(mode=CleanMode.ALL))
        self.assertTrue(result.success, result.error or result.residual_items)
        assert result.output is not None
        with ZipFile(result.output) as zf:
            names = set(zf.namelist())
            self.assertNotIn("docProps/core.xml", names)
            self.assertNotIn("docProps/app.xml", names)
            self.assertNotIn("docProps/custom.xml", names)
            self.assertNotIn("docProps/thumbnail.jpeg", names)
        self.assertEqual(Document(result.output).paragraphs[0].text, "正文内容不得改变")

    def test_signed_docx_is_skipped_without_output(self) -> None:
        source = self.root / "signed.docx"
        make_docx(source, signed=True)
        result = clean_file(source)
        self.assertFalse(result.success)
        self.assertIsNone(result.output)
        self.assertIn("数字签名", result.error)
        self.assertEqual(list(self.root.glob("signed_clean*.docx")), [])

    def test_user_can_intentionally_keep_thumbnail_and_reviewer_metadata(self) -> None:
        source = self.root / "keep-options.docx"
        make_docx(source)
        options = CleanOptions(
            mode=CleanMode.PERSONAL,
            anonymize_reviewers=False,
            remove_thumbnail=False,
            verify_after_clean=True,
        )
        result = clean_file(source, options=options)
        self.assertTrue(result.success, result.error or result.residual_items)
        assert result.output is not None
        with ZipFile(result.output) as zf:
            self.assertIn("docProps/thumbnail.jpeg", zf.namelist())
            self.assertIn("Reviewer Name", zf.read("word/comments.xml").decode("utf-8"))
            self.assertIn("Header Reviewer", zf.read("word/header1.xml").decode("utf-8"))

    def test_corrupt_office_file_returns_a_failure_result(self) -> None:
        source = self.root / "corrupt.docx"
        source.write_bytes(b"this is not a zip package")
        inspection = inspect_file(source)
        self.assertFalse(inspection.supported)
        result = clean_file(source)
        self.assertFalse(result.success)
        self.assertIsNone(result.output)
        self.assertIn("不支持", result.error)

    def test_pdf_personal_and_all_modes(self) -> None:
        source = self.root / "sample.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "CONTENT-CHECK")
        doc.set_metadata(
            {
                "author": "Alice Example",
                "creator": "Sensitive Workstation",
                "producer": "Sensitive Producer",
                "title": "保留的 PDF 标题",
                "subject": "保留的主题",
                "keywords": "uranium,membrane",
            }
        )
        doc.save(source)
        doc.close()

        personal = clean_file(source, options=CleanOptions(mode=CleanMode.PERSONAL))
        self.assertTrue(personal.success, personal.error or personal.residual_items)
        assert personal.output is not None
        with fitz.open(personal.output) as cleaned:
            self.assertIn("CONTENT-CHECK", cleaned[0].get_text())
            self.assertEqual(cleaned.metadata.get("author"), "")
            self.assertEqual(cleaned.metadata.get("creator"), "")
            self.assertEqual(cleaned.metadata.get("title"), "保留的 PDF 标题")

        all_result = clean_file(source, options=CleanOptions(mode=CleanMode.ALL))
        self.assertTrue(all_result.success, all_result.error or all_result.residual_items)
        assert all_result.output is not None
        with fitz.open(all_result.output) as cleaned:
            self.assertIn("CONTENT-CHECK", cleaned[0].get_text())
            self.assertEqual(cleaned.metadata.get("title"), "")
            self.assertEqual(cleaned.metadata.get("subject"), "")

    def test_encrypted_pdf_is_skipped(self) -> None:
        source = self.root / "encrypted.pdf"
        doc = fitz.open()
        doc.new_page().insert_text((72, 72), "SECRET")
        doc.save(
            source,
            encryption=fitz.PDF_ENCRYPT_AES_256,
            owner_pw="owner-password",
            user_pw="user-password",
        )
        doc.close()
        result = clean_file(source)
        self.assertFalse(result.success)
        self.assertIsNone(result.output)
        self.assertIn("加密", result.error)

    def test_pdf_custom_info_xmp_and_annotation_identity_are_cleaned(self) -> None:
        source = self.root / "rich-metadata.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "PDF-BODY")
        doc.set_metadata({"author": "Info Author", "title": "保留 Info 标题"})
        info_xref = int(doc.xref_get_key(-1, "Info")[1].split()[0])
        doc.xref_set_key(info_xref, "Company", "(SecretCorp)")
        annotation = page.add_text_annot((100, 100), "批注正文保留")
        annotation.set_info(
            title="Annotation Reviewer",
            creationDate="D:20260721010203+08'00'",
            modDate="D:20260721020304+08'00'",
            subject="保留批注主题",
        )
        xmp = """<?xpacket begin="﻿" id="W5M0MpCehiHzreSzNTczkc9d"?>
<x:xmpmeta xmlns:x="adobe:ns:meta/">
 <rdf:RDF xmlns:rdf="http://www.w3.org/1999/02/22-rdf-syntax-ns#">
  <rdf:Description rdf:about=""
   xmlns:dc="http://purl.org/dc/elements/1.1/"
   xmlns:xmp="http://ns.adobe.com/xap/1.0/"
    xmlns:pdfaid="http://www.aiim.org/pdfa/ns/id/"
    xmlns:xmpMM="http://ns.adobe.com/xap/1.0/mm/"
    xmlns:stEvt="http://ns.adobe.com/xap/1.0/sType/ResourceEvent#"
    xmlns:Iptc4xmpCore="http://iptc.org/std/Iptc4xmpCore/1.0/xmlns/"
    xmlns:evil="https://example.test/evil/">
   <dc:title><rdf:Alt><rdf:li xml:lang="x-default">保留 XMP 标题</rdf:li></rdf:Alt></dc:title>
   <dc:creator><rdf:Seq><rdf:li>XMP Author</rdf:li></rdf:Seq></dc:creator>
   <xmp:CreatorTool>Sensitive Tool</xmp:CreatorTool>
    <xmp:CreateDate>2026-07-21T01:02:03+08:00</xmp:CreateDate>
    <xmpMM:History><rdf:Seq><rdf:li rdf:parseType="Resource"><stEvt:softwareAgent>Private Editor</stEvt:softwareAgent><stEvt:when>2026-07-21T03:04:05+08:00</stEvt:when></rdf:li></rdf:Seq></xmpMM:History>
    <Iptc4xmpCore:CreatorContactInfo><Iptc4xmpCore:CiEmailWork>author@example.test</Iptc4xmpCore:CiEmailWork></Iptc4xmpCore:CreatorContactInfo>
    <evil:title>Namespace Bypass Secret</evil:title>
    <evil:part>Namespace Part Secret</evil:part>
   <pdfaid:part>2</pdfaid:part>
  </rdf:Description>
 </rdf:RDF>
</x:xmpmeta>
<?xpacket end="w"?>"""
        doc.set_xml_metadata(xmp)
        doc.save(source)
        doc.close()

        before = inspect_file(source)
        self.assertTrue(any(item.group == "PDF 自定义属性" for item in before.items))
        self.assertTrue(any(item.group == "PDF 批注属性" for item in before.items))
        self.assertTrue(any(item.group == "PDF XMP" and item.name == "creator" for item in before.items))
        self.assertTrue(any(item.group == "PDF XMP" and item.name == "History" for item in before.items))

        result = clean_file(source, options=CleanOptions(mode=CleanMode.PERSONAL))
        self.assertTrue(result.success, result.error or result.residual_items)
        assert result.output is not None
        with fitz.open(result.output) as cleaned:
            self.assertIn("PDF-BODY", cleaned[0].get_text())
            xref = int(cleaned.xref_get_key(-1, "Info")[1].split()[0])
            self.assertEqual(cleaned.xref_get_key(xref, "Company")[0], "null")
            xmp_after = cleaned.get_xml_metadata()
            self.assertIn("保留 XMP 标题", xmp_after)
            self.assertIn("http://www.aiim.org/pdfa/ns/id/", xmp_after)
            self.assertIn(">2</", xmp_after)
            self.assertNotIn("XMP Author", xmp_after)
            self.assertNotIn("Sensitive Tool", xmp_after)
            self.assertNotIn("Private Editor", xmp_after)
            self.assertNotIn("2026-07-21T03:04:05+08:00", xmp_after)
            self.assertNotIn("author@example.test", xmp_after)
            self.assertNotIn("Namespace Bypass Secret", xmp_after)
            self.assertNotIn("Namespace Part Secret", xmp_after)
            self.assertEqual(cleaned.xref_get_key(-1, "ID")[0], "null")
            annot_after = next(cleaned[0].annots())
            self.assertEqual(annot_after.info.get("content"), "批注正文保留")
            self.assertEqual(annot_after.info.get("subject"), "保留批注主题")
            self.assertEqual(annot_after.info.get("title"), "")
            self.assertEqual(annot_after.info.get("creationDate"), "")
            self.assertEqual(annot_after.info.get("modDate"), "")
        self.assertEqual(inspect_file(result.output).personal_count, 0)

    def test_odt_personal_mode_preserves_title_and_body(self) -> None:
        source = self.root / "sample.odt"
        make_odt(source)
        result = clean_file(source, options=CleanOptions(mode=CleanMode.PERSONAL))
        self.assertTrue(result.success, result.error or result.residual_items)
        assert result.output is not None
        with ZipFile(result.output) as zf:
            self.assertEqual(zf.infolist()[0].filename, "mimetype")
            self.assertEqual(zf.infolist()[0].compress_type, ZIP_STORED)
            meta = zf.read("meta.xml").decode("utf-8")
            content = zf.read("content.xml").decode("utf-8")
            self.assertIn("保留的 ODF 标题", meta)
            self.assertNotIn("Alice Example", meta)
            self.assertNotIn("Bob Example", meta)
            self.assertNotIn("EMP-001", meta)
            self.assertNotIn("private-template.ott", meta)
            self.assertIn("正文内容不得改变", content)
            self.assertIn("2026-08-15", content)
            self.assertNotIn("Reviewer Name", content)
            self.assertNotIn("2026-07-21", content)
        self.assertEqual(inspect_file(result.output).personal_count, 0)

    def test_jpeg_and_png_metadata_are_removed(self) -> None:
        jpeg = self.root / "photo.jpg"
        image = Image.new("RGB", (40, 30), (12, 100, 200))
        exif = Image.Exif()
        exif[315] = "Alice Example"
        exif[305] = "Sensitive Software"
        image.save(jpeg, quality=92, exif=exif)
        jpeg_result = clean_file(jpeg)
        self.assertTrue(jpeg_result.success, jpeg_result.error or jpeg_result.residual_items)
        assert jpeg_result.output is not None
        with Image.open(jpeg_result.output) as cleaned:
            self.assertEqual(cleaned.size, (40, 30))
            self.assertEqual(len(cleaned.getexif()), 0)

        png = self.root / "figure.png"
        pnginfo = PngInfo()
        pnginfo.add_text("Author", "Alice Example")
        pnginfo.add_text("Description", "Sensitive note")
        image.save(png, pnginfo=pnginfo)
        png_result = clean_file(png)
        self.assertTrue(png_result.success, png_result.error or png_result.residual_items)
        assert png_result.output is not None
        with Image.open(png_result.output) as cleaned:
            self.assertNotIn("Author", cleaned.info)
            self.assertNotIn("Description", cleaned.info)

    def test_all_mode_removes_png_icc_profile(self) -> None:
        source = self.root / "profile.png"
        marker = b"PRIVATE-ICC-PROFILE-Alice-Workstation"
        Image.new("RGB", (8, 8), (10, 20, 30)).save(source, icc_profile=marker)

        before = inspect_file(source)
        self.assertTrue(any(item.name == "ICC profile" for item in before.items))
        result = clean_file(source, options=CleanOptions(mode=CleanMode.ALL))
        self.assertTrue(result.success, result.error or result.residual_items)
        assert result.output is not None
        with Image.open(result.output) as cleaned:
            self.assertNotIn("icc_profile", cleaned.info)
        self.assertNotIn(marker, result.output.read_bytes())

    def test_tiff_icc_is_technical_in_personal_mode_and_removed_in_all_mode(self) -> None:
        source = self.root / "profile.tiff"
        profile = ImageCms.ImageCmsProfile(ImageCms.createProfile("sRGB")).tobytes()
        Image.new("RGB", (8, 8), (10, 20, 30)).save(source, icc_profile=profile)

        personal = clean_file(source, options=CleanOptions(mode=CleanMode.PERSONAL))
        self.assertTrue(personal.success, personal.error or personal.residual_items)
        assert personal.output is not None
        with Image.open(personal.output) as cleaned:
            self.assertIn("icc_profile", cleaned.info)

        all_result = clean_file(source, options=CleanOptions(mode=CleanMode.ALL))
        self.assertTrue(all_result.success, all_result.error or all_result.residual_items)
        assert all_result.output is not None
        with Image.open(all_result.output) as cleaned:
            self.assertNotIn("icc_profile", cleaned.info)

    def test_owner_password_only_pdf_is_still_treated_as_encrypted(self) -> None:
        source = self.root / "owner-only.pdf"
        doc = fitz.open()
        doc.new_page().insert_text((72, 72), "OWNER-PROTECTED")
        doc.save(
            source,
            encryption=fitz.PDF_ENCRYPT_AES_256,
            owner_pw="owner-password",
            user_pw="",
        )
        doc.close()

        inspection = inspect_file(source)
        self.assertTrue(inspection.encrypted)
        result = clean_file(source)
        self.assertFalse(result.success)
        self.assertIsNone(result.output)
        self.assertIn("加密", result.error)

    def test_encrypted_odf_member_is_skipped(self) -> None:
        source = self.root / "encrypted.odt"
        make_odt(source)
        with ZipFile(source) as zf:
            manifest = ET.fromstring(zf.read("META-INF/manifest.xml"))
        file_entry = next(elem for elem in manifest if elem.tag.endswith("file-entry"))
        ET.SubElement(
            file_entry,
            "{urn:oasis:names:tc:opendocument:xmlns:manifest:1.0}encryption-data",
        )
        rewrite_zip(
            source,
            {"META-INF/manifest.xml": ET.tostring(manifest, encoding="utf-8", xml_declaration=True)},
        )

        inspection = inspect_file(source)
        self.assertTrue(inspection.encrypted)
        result = clean_file(source)
        self.assertFalse(result.success)
        self.assertIsNone(result.output)
        self.assertIn("加密", result.error)

    def test_malformed_odf_manifest_fails_closed(self) -> None:
        source = self.root / "malformed-manifest.odt"
        make_odt(source)
        rewrite_zip(
            source,
            {"META-INF/manifest.xml": b"<manifest:manifest><manifest:encryption-data>"},
        )

        inspection = inspect_file(source)
        self.assertFalse(inspection.supported)
        result = clean_file(source)
        self.assertFalse(result.success)
        self.assertIsNone(result.output)

    def test_pdf_with_embedded_attachment_requires_manual_review(self) -> None:
        source = self.root / "attachment.pdf"
        doc = fitz.open()
        doc.new_page().insert_text((72, 72), "BODY")
        doc.embfile_add("private-notes.txt", b"embedded content", filename="private-notes.txt")
        doc.save(source)
        doc.close()

        before = inspect_file(source)
        self.assertTrue(any(item.group == "PDF 嵌入附件" for item in before.items))
        result = clean_file(source)
        self.assertFalse(result.success)
        self.assertIsNotNone(result.output)
        self.assertTrue(any(item.group == "PDF 嵌入附件" for item in result.residual_items))

    def test_output_race_never_overwrites_or_deletes_external_file(self) -> None:
        sources: list[Path] = []

        docx = self.root / "race.docx"
        make_docx(docx)
        sources.append(docx)

        odt = self.root / "race.odt"
        make_odt(odt)
        sources.append(odt)

        pdf = self.root / "race.pdf"
        doc = fitz.open()
        doc.new_page().insert_text((72, 72), "RACE")
        doc.save(pdf)
        doc.close()
        sources.append(pdf)

        png = self.root / "race.png"
        Image.new("RGB", (8, 8), (1, 2, 3)).save(png)
        sources.append(png)

        sentinel = b"created by another process"
        for index, source in enumerate(sources):
            raced_output = self.root / f"external-{index}{source.suffix}"

            def reserve_raced_path(*_args: object, path: Path = raced_output, **_kwargs: object) -> Path:
                path.write_bytes(sentinel)
                return path

            with self.subTest(suffix=source.suffix), patch(
                "metacleaner.engine.unique_output_path",
                side_effect=reserve_raced_path,
            ):
                result = clean_file(source)
                self.assertFalse(result.success)
                self.assertIsNone(result.output)
                self.assertEqual(raced_output.read_bytes(), sentinel)

    def test_unique_output_path_never_overwrites(self) -> None:
        source = self.root / "report.pdf"
        source.write_bytes(b"source")
        first = unique_output_path(source)
        self.assertEqual(first.name, "report_clean.pdf")
        first.write_bytes(b"first")
        second = unique_output_path(source)
        self.assertEqual(second.name, "report_clean_2.pdf")

    def test_default_namespace_serialization_keeps_qualified_attributes_bound(self) -> None:
        namespace = "http://schemas.openxmlformats.org/package/2006/relationships"
        root = ET.fromstring(
            f'<rel:Relationships xmlns:rel="{namespace}">'
            '<rel:Relationship rel:Id="rId1" Target="word/document.xml"/>'
            '</rel:Relationships>'
        )
        serialized = _serialize_with_default_namespace(root)
        self.assertIn(f'xmlns="{namespace}"'.encode(), serialized)
        self.assertIn(f'xmlns:rel="{namespace}"'.encode(), serialized)
        self.assertNotIn(b"<rel:Relationships", serialized)
        parsed = ET.fromstring(serialized)
        relationship = next(iter(parsed))
        self.assertEqual(relationship.attrib[f"{{{namespace}}}Id"], "rId1")


if __name__ == "__main__":
    unittest.main()
